"""文档工作台的存储与本地提取编排。"""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
import json
from pathlib import Path
import re
import time
from typing import Literal
from uuid import UUID

from docx import Document as WordDocument

from iris_agent.core.models import ProviderResponse
from iris_agent.providers.base import ModelProvider
from iris_agent.documents.drafts import DraftRepository
from iris_agent.documents.errors import (
    DocumentError,
    DocumentExtractFailedError,
    DocumentGenerationError,
    DocumentValidationError,
)
from iris_agent.documents.extraction import LocalDocumentExtractor
from iris_agent.documents.models import DOCUMENT_TEMPLATES, DocumentCitation, DocumentDraft, DocumentRecord
from iris_agent.documents.prompts import build_draft_messages
from iris_agent.documents.storage import DocumentRepository


DocumentExportFormat = Literal["markdown", "docx"]


@dataclass(frozen=True, slots=True)
class DocumentExport:
    content: bytes
    media_type: str
    filename: str


@dataclass(slots=True)
class DocumentService:
    root: Path
    provider: ModelProvider | None = None
    max_file_bytes: int = 10_000_000
    max_total_bytes: int = 50_000_000
    max_count: int = 50
    max_text_chars: int = 50_000
    max_draft_chars: int = 100_000
    repository: DocumentRepository = field(init=False)
    extractor: LocalDocumentExtractor = field(init=False)
    drafts: DraftRepository = field(init=False)

    def __post_init__(self) -> None:
        self.repository = DocumentRepository(
            self.root,
            max_file_bytes=self.max_file_bytes,
            max_total_bytes=self.max_total_bytes,
            max_count=self.max_count,
            max_text_chars=self.max_text_chars,
        )
        self.extractor = LocalDocumentExtractor(self.max_text_chars)
        if not isinstance(self.max_draft_chars, int) or isinstance(self.max_draft_chars, bool) or self.max_draft_chars < 1:
            raise DocumentValidationError("文档草稿长度限制必须大于 0")
        self.drafts = DraftRepository(self.repository.root)

    def upload(self, original_name: str, content: bytes, media_type: str) -> DocumentRecord:
        document = self.repository.save(original_name, content, media_type)
        try:
            extraction = self.extractor.extract(self.repository.file_for(document.id))
        except DocumentExtractFailedError:
            return self.repository.update_extraction(document.id, None, message="无法提取文档文本")
        return self.repository.update_extraction(document.id, extraction)

    def list(self) -> list[DocumentRecord]:
        return self.repository.list()

    def get(self, document_id: str) -> DocumentRecord:
        return self.repository.get(document_id)

    def read_text(self, document_id: str) -> str:
        return self.repository.read_text(document_id)

    def delete(self, document_id: str) -> None:
        self.repository.delete(document_id)

    def generate_draft(
        self,
        template: str,
        document_ids: list[str] | tuple[str, ...],
        instructions: str = "",
    ) -> DocumentDraft:
        selected_ids = self._validate_sources(document_ids)
        safe_instructions = self._validate_instructions(instructions)
        if template not in DOCUMENT_TEMPLATES:
            raise DocumentValidationError("不支持该文档模板")
        source_documents: list[dict[str, object]] = []
        for document_id in selected_ids:
            document = self.repository.get(document_id)
            if document.extraction_status != "ready":
                raise DocumentValidationError("所选资料尚未准备好")
            source_documents.append(
                {
                    "id": document.id,
                    "name": document.original_name,
                    "text": self.repository.read_text(document.id),
                    "truncated": document.text_truncated,
                }
            )
        if self.provider is None:
            raise DocumentGenerationError("文档生成服务暂不可用")
        try:
            response = self.provider.complete(
                build_draft_messages(template, source_documents, safe_instructions),
                [],
            )
        except Exception as exc:
            raise DocumentGenerationError("生成文档草稿失败，请稍后重试") from exc
        title, markdown, citations = self._parse_generated_draft(response, set(selected_ids))
        if len(markdown) > self.max_draft_chars:
            raise DocumentGenerationError("模型返回的文档草稿过长", code="document_model_output_invalid")
        return self.drafts.create(
            title=title,
            template=template,
            document_ids=selected_ids,
            instructions=safe_instructions,
            markdown=markdown,
            citations=citations,
        )

    def list_drafts(self) -> list[DocumentDraft]:
        return self.drafts.list()

    def get_draft(self, draft_id: str) -> DocumentDraft:
        return self.drafts.get(draft_id)

    def save_draft(
        self,
        draft_id: str,
        title: str,
        markdown: str,
        expected_revision: int,
    ) -> DocumentDraft:
        safe_markdown = self._validate_markdown(markdown)
        if len(safe_markdown) > self.max_draft_chars:
            raise DocumentValidationError("文档草稿超过长度限制")
        return self.drafts.update(
            draft_id,
            title=self._validate_title(title),
            markdown=safe_markdown,
            expected_revision=expected_revision,
        )

    def export_draft(self, draft_id: str, format: str) -> DocumentExport:
        draft = self.get_draft(draft_id)
        if format == "markdown":
            return DocumentExport(
                content=self._render_markdown(draft).encode("utf-8"),
                media_type="text/markdown",
                filename=f"iris-document-{draft.id}.md",
            )
        if format == "docx":
            return DocumentExport(
                content=self._render_docx(draft),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"iris-document-{draft.id}.docx",
            )
        raise DocumentValidationError("不支持该导出格式")

    @staticmethod
    def _validate_sources(document_ids: list[str] | tuple[str, ...]) -> tuple[str, ...]:
        if not isinstance(document_ids, (list, tuple)) or not document_ids:
            raise DocumentValidationError("至少选择一份已准备好的资料")
        if any(not isinstance(item, str) for item in document_ids) or len(document_ids) != len(set(document_ids)):
            raise DocumentValidationError("所选资料无效")
        for document_id in document_ids:
            try:
                if str(UUID(document_id)) != document_id:
                    raise ValueError
            except (TypeError, ValueError, AttributeError):
                raise DocumentValidationError("所选资料无效") from None
        return tuple(document_ids)

    @staticmethod
    def _validate_instructions(instructions: str) -> str:
        if not isinstance(instructions, str) or len(instructions) > 2_000:
            raise DocumentValidationError("文档生成说明无效")
        return instructions.strip()

    @staticmethod
    def _validate_title(title: str) -> str:
        if (
            not isinstance(title, str)
            or not title.strip()
            or len(title.strip()) > 200
            or "\r" in title
            or "\n" in title
        ):
            raise DocumentValidationError("文档草稿标题无效")
        return title.strip()

    @staticmethod
    def _validate_markdown(markdown: str) -> str:
        if not isinstance(markdown, str) or not markdown.strip():
            raise DocumentValidationError("文档草稿正文无效")
        return markdown.strip()

    @classmethod
    def _parse_generated_draft(
        cls,
        response: ProviderResponse,
        allowed_document_ids: set[str],
    ) -> tuple[str, str, tuple[DocumentCitation, ...]]:
        content = response.content.strip()
        try:
            if content.startswith("```") or content.endswith("```"):
                raise ValueError("markdown fence")
            raw = json.loads(content)
            if not isinstance(raw, dict) or set(raw) != {"title", "markdown", "citations"}:
                raise ValueError("unexpected keys")
            title = cls._validate_title(raw["title"])
            markdown = cls._validate_markdown(raw["markdown"])
            citations = raw["citations"]
            if not isinstance(citations, list):
                raise ValueError("invalid citations")
            parsed = tuple(
                DocumentCitation(document_id=item["document_id"], location=item["location"])
                for item in citations
                if isinstance(item, dict) and set(item) == {"document_id", "location"}
            )
            if len(parsed) != len(citations) or any(
                citation.document_id not in allowed_document_ids for citation in parsed
            ):
                raise ValueError("untrusted citations")
            return title, markdown, parsed
        except (json.JSONDecodeError, KeyError, TypeError, ValueError, DocumentError) as exc:
            raise DocumentGenerationError(
                "模型返回的文档草稿格式无效",
                code="document_model_output_invalid",
            ) from exc

    @staticmethod
    def _render_markdown(draft: DocumentDraft) -> str:
        sources = "\n".join(
            f"- {citation.document_id}: {citation.location}" for citation in draft.citations
        )
        blocks = [f"# {draft.title}", draft.markdown]
        if sources:
            blocks.append(f"## Sources\n{sources}")
        return "\n\n".join(blocks) + "\n"

    @classmethod
    def _render_docx(cls, draft: DocumentDraft) -> bytes:
        document = WordDocument()
        document.add_heading(draft.title, level=0)
        for line in draft.markdown.splitlines():
            value = line.strip()
            if not value:
                continue
            heading = re.fullmatch(r"(#{1,6})\s+(.+)", value)
            if heading is not None:
                document.add_heading(heading.group(2), level=min(len(heading.group(1)), 6))
            elif value.startswith("- ") or value.startswith("* "):
                document.add_paragraph(value[2:], style="List Bullet")
            else:
                document.add_paragraph(value)
        if draft.citations:
            document.add_heading("Sources", level=1)
            for citation in draft.citations:
                document.add_paragraph(f"{citation.document_id}: {citation.location}", style="List Bullet")
        output = BytesIO()
        document.save(output)
        return output.getvalue()
