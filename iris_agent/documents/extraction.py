"""本地文档文本提取，不依赖日报附件实现。"""

from __future__ import annotations

from io import BytesIO
from typing import Protocol

from iris_agent.documents.errors import DocumentExtractFailedError, DocumentInvalidTypeError
from iris_agent.documents.models import DocumentExtraction, DocumentSource


_TEXT_SUFFIXES = frozenset({".txt", ".md"})
_SUPPORTED_SUFFIXES = _TEXT_SUFFIXES | frozenset({".docx", ".pdf", ".xlsx", ".xls"})


class ExtractableDocument(Protocol):
    name: str
    suffix: str

    def read_bytes(self) -> bytes: ...


class LocalDocumentExtractor:
    """使用已安装的本地解析器提取有限长度的文档正文。"""

    def __init__(self, max_text_chars: int):
        if not isinstance(max_text_chars, int) or isinstance(max_text_chars, bool) or max_text_chars < 1:
            raise ValueError("max_text_chars must be positive")
        self._max_text_chars = max_text_chars

    def extract(self, document: ExtractableDocument) -> DocumentExtraction:
        suffix = document.suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES:
            raise DocumentExtractFailedError("不支持提取该文档类型")
        try:
            content = document.read_bytes()
            text, locations = self._extract_content(suffix, content)
        except DocumentInvalidTypeError:
            raise
        except Exception as exc:
            raise DocumentExtractFailedError("无法提取文档文本") from exc
        if not isinstance(text, str) or not text.strip():
            raise DocumentExtractFailedError("文档未包含可提取文本")
        limited = text[: self._max_text_chars]
        sources = tuple(DocumentSource(document.name, location) for location in locations)
        if not sources:
            sources = (DocumentSource(document.name, "正文"),)
        return DocumentExtraction(text=limited, sources=sources, truncated=len(text) > len(limited))

    def _extract_content(self, suffix: str, content: bytes) -> tuple[str, tuple[str, ...]]:
        if not isinstance(content, bytes):
            raise DocumentInvalidTypeError("文档内容无效")
        if suffix in _TEXT_SUFFIXES:
            return content.decode("utf-8"), ("正文",)
        if suffix == ".docx":
            return self._extract_docx(content)
        if suffix == ".pdf":
            return self._extract_pdf(content)
        if suffix == ".xlsx":
            return self._extract_xlsx(content)
        return self._extract_xls(content)

    @staticmethod
    def _extract_docx(content: bytes) -> tuple[str, tuple[str, ...]]:
        from docx import Document

        document = Document(BytesIO(content))
        parts: list[str] = []
        locations: list[str] = []
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        if paragraphs:
            parts.extend(paragraphs)
            locations.append("段落")
        for number, table in enumerate(document.tables, start=1):
            rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            rows = [row for row in rows if row.strip()]
            if rows:
                parts.extend(rows)
                locations.append(f"表格 {number}")
        return "\n".join(parts), tuple(locations)

    @staticmethod
    def _extract_pdf(content: bytes) -> tuple[str, tuple[str, ...]]:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        parts: list[str] = []
        locations: list[str] = []
        for number, page in enumerate(reader.pages, start=1):
            extracted = (page.extract_text() or "").strip()
            if extracted:
                parts.append(extracted)
                locations.append(f"第 {number} 页")
        return "\n".join(parts), tuple(locations)

    @staticmethod
    def _extract_xlsx(content: bytes) -> tuple[str, tuple[str, ...]]:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        try:
            parts: list[str] = []
            locations: list[str] = []
            for worksheet in workbook.worksheets:
                location = f"工作表：{worksheet.title}"
                parts.append(location)
                locations.append(location)
                for row in worksheet.iter_rows(min_row=1, max_row=100, max_col=20, values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(values):
                        parts.append("\t".join(values))
            return "\n".join(parts), tuple(locations)
        finally:
            workbook.close()

    @staticmethod
    def _extract_xls(content: bytes) -> tuple[str, tuple[str, ...]]:
        import xlrd

        workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
        try:
            parts: list[str] = []
            locations: list[str] = []
            for worksheet in workbook.sheets():
                location = f"工作表：{worksheet.name}"
                parts.append(location)
                locations.append(location)
                for row_index in range(min(worksheet.nrows, 100)):
                    values = [
                        "" if value is None else str(value)
                        for value in (
                            worksheet.cell_value(row_index, column_index)
                            for column_index in range(min(worksheet.ncols, 20))
                        )
                    ]
                    if any(values):
                        parts.append("\t".join(values))
            return "\n".join(parts), tuple(locations)
        finally:
            release = getattr(workbook, "release_resources", None)
            if callable(release):
                release()
