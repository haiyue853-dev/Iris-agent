from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Protocol

from iris_agent.reports.attachments import AttachmentFile
from iris_agent.reports.errors import (
    ReportAttachmentExtractError,
    ReportAttachmentInvalidTypeError,
    ReportAttachmentOcrUnavailableError,
)


_TEXT_SUFFIXES = frozenset({".txt", ".md"})
_IMAGE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".webp"})
_SUPPORTED_SUFFIXES = _TEXT_SUFFIXES | _IMAGE_SUFFIXES | frozenset({".docx", ".pdf", ".xlsx", ".xls"})


@dataclass(frozen=True, slots=True)
class AttachmentExtractionSource:
    file_name: str
    location: str | None = None


@dataclass(frozen=True, slots=True)
class AttachmentExtraction:
    text: str
    sources: tuple[AttachmentExtractionSource, ...]
    truncated: bool


class LocalOcr(Protocol):
    def extract(self, content: bytes, file_name: str) -> str: ...


class LocalAttachmentExtractor:
    def __init__(self, max_chars: int, ocr: LocalOcr | None = None):
        if max_chars < 1:
            raise ValueError("max_chars must be positive")
        self._max_chars = max_chars
        self._ocr = ocr

    def extract(self, attachment: AttachmentFile) -> AttachmentExtraction:
        suffix = attachment.suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES:
            raise ReportAttachmentInvalidTypeError("不支持该日报附件")
        try:
            content = attachment.read_bytes()
            text, locations = self._extract_content(suffix, content, attachment.name)
        except ReportAttachmentOcrUnavailableError:
            raise
        except Exception as exc:
            raise ReportAttachmentExtractError("无法提取日报附件文本") from exc
        if not isinstance(text, str) or not text.strip():
            raise ReportAttachmentExtractError("日报附件未包含可提取文本")
        limited = text[: self._max_chars]
        return AttachmentExtraction(
            text=limited,
            sources=tuple(AttachmentExtractionSource(attachment.name, location) for location in locations) or (AttachmentExtractionSource(attachment.name),),
            truncated=len(text) > len(limited),
        )

    def _extract_content(self, suffix: str, content: bytes, file_name: str) -> tuple[str, tuple[str | None, ...]]:
        if suffix in _TEXT_SUFFIXES:
            return content.decode("utf-8"), (None,)
        if suffix == ".docx":
            return self._extract_docx(content)
        if suffix == ".pdf":
            return self._extract_pdf(content)
        if suffix == ".xlsx":
            return self._extract_excel(content)
        if suffix == ".xls":
            return self._extract_xls(content)
        if self._ocr is None:
            raise ReportAttachmentOcrUnavailableError("本机 OCR 未配置")
        return self._ocr.extract(content, file_name), (None,)

    @staticmethod
    def _extract_docx(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        from docx import Document

        document = Document(BytesIO(content))
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            lines.extend("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return "\n".join(lines), (None,)

    @staticmethod
    def _extract_pdf(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n".join(page for page in pages if page), tuple(f"第 {number} 页" for number, page in enumerate(pages, start=1) if page)

    @staticmethod
    def _extract_excel(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        try:
            parts: list[str] = []
            locations: list[str] = []
            for worksheet in workbook.worksheets:
                parts.append(f"工作表：{worksheet.title}")
                locations.append(f"工作表：{worksheet.title}")
                for row in worksheet.iter_rows(min_row=1, max_row=100, max_col=20, values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(values):
                        parts.append("\t".join(values))
            return "\n".join(parts), tuple(locations)
        finally:
            workbook.close()

    @staticmethod
    def _extract_xls(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        import xlrd

        workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
        try:
            parts: list[str] = []
            locations: list[str] = []
            for worksheet in workbook.sheets():
                parts.append(f"工作表：{worksheet.name}")
                locations.append(f"工作表：{worksheet.name}")
                for row_index in range(min(worksheet.nrows, 100)):
                    values = [LocalAttachmentExtractor._cell_text(worksheet.cell_value(row_index, column_index)) for column_index in range(min(worksheet.ncols, 20))]
                    if any(values):
                        parts.append("\t".join(values))
            return "\n".join(parts), tuple(locations)
        finally:
            close = getattr(workbook, "release_resources", None)
            if callable(close):
                close()

    @staticmethod
    def _cell_text(value: object) -> str:
        return "" if value is None else str(value)

