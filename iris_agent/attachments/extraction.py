from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Protocol

from .errors import AttachmentExtractError, AttachmentInvalidTypeError
from .storage import AttachmentFile

_TEXT_SUFFIXES = frozenset({".txt", ".md"})
_IMAGE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".webp"})
_SUPPORTED_SUFFIXES = _TEXT_SUFFIXES | _IMAGE_SUFFIXES | frozenset({".docx", ".pdf", ".xlsx", ".xls"})

@dataclass(frozen=True, slots=True)
class AttachmentExtractionSource:
    file_name: str
    location: str | None = None

@dataclass(frozen=True, slots=True)
class AttachmentExtraction:
    text: str
    sources: tuple[AttachmentExtractionSource, ...]
    truncated: bool

class LocalOcr(Protocol):
    def extract(self, content: bytes, file_name: str) -> str: ...

class LocalAttachmentExtractor:
    def __init__(self, max_chars: int, ocr: LocalOcr | None = None):
        if max_chars < 1: raise ValueError("max_chars must be positive")
        self._max_chars, self._ocr = max_chars, ocr

    def extract(self, attachment: AttachmentFile) -> AttachmentExtraction:
        suffix = attachment.suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES: raise AttachmentInvalidTypeError("不支持该附件")
        original_name = getattr(attachment, "original_name", attachment.name)
        try:
            text, locations = self._extract_content(suffix, attachment.read_bytes(), original_name)
        except AttachmentExtractError: raise
        except Exception as exc: raise AttachmentExtractError("无法提取附件文本") from exc
        if not text.strip(): raise AttachmentExtractError("附件未包含可提取文本")
        limited = text[:self._max_chars]
        return AttachmentExtraction(limited, tuple(AttachmentExtractionSource(original_name, location) for location in locations) or (AttachmentExtractionSource(original_name),), len(text) > len(limited))

    def _extract_content(self, suffix: str, content: bytes, name: str) -> tuple[str, tuple[str | None, ...]]:
        if suffix in _TEXT_SUFFIXES: return content.decode("utf-8"), (None,)
        if suffix == ".docx": return self._extract_docx(content)
        if suffix == ".pdf": return self._extract_pdf(content)
        if suffix == ".xlsx": return self._extract_excel(content)
        if suffix == ".xls": return self._extract_xls(content)
        if self._ocr is None: raise AttachmentExtractError("本机 OCR 未配置")
        return self._ocr.extract(content, name), (None,)

    @staticmethod
    def _extract_docx(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        from docx import Document
        document = Document(BytesIO(content)); lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables: lines.extend("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return "\n".join(lines), (None,)

    @staticmethod
    def _extract_pdf(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        from pypdf import PdfReader
        pages = [(page.extract_text() or "").strip() for page in PdfReader(BytesIO(content)).pages]
        return "\n".join(page for page in pages if page), tuple(f"第 {index} 页" for index, page in enumerate(pages, 1) if page)

    @staticmethod
    def _extract_excel(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        from openpyxl import load_workbook
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        try:
            parts: list[str] = []; locations: list[str] = []
            for sheet in workbook.worksheets:
                location = f"工作表：{sheet.title}"; parts.append(location); locations.append(location)
                for row in sheet.iter_rows(min_row=1, max_row=100, max_col=20, values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(values): parts.append("\t".join(values))
            return "\n".join(parts), tuple(locations)
        finally: workbook.close()

    @staticmethod
    def _extract_xls(content: bytes) -> tuple[str, tuple[str | None, ...]]:
        import xlrd
        workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
        try:
            parts: list[str] = []; locations: list[str] = []
            for sheet in workbook.sheets():
                location = f"工作表：{sheet.name}"; parts.append(location); locations.append(location)
                for row in range(min(sheet.nrows, 100)):
                    values = ["" if (value := sheet.cell_value(row, column)) is None else str(value) for column in range(min(sheet.ncols, 20))]
                    if any(values): parts.append("\t".join(values))
            return "\n".join(parts), tuple(locations)
        finally:
            close = getattr(workbook, "release_resources", None)
            if callable(close): close()
