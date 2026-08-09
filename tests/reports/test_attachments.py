from __future__ import annotations

import json
import math
import os
from pathlib import Path
import shutil
import threading

import pytest

from iris_agent.reports import (
    ReportAttachmentError,
    ReportAttachmentInvalidTypeError,
    ReportAttachmentNotFoundError,
    ReportAttachmentStorageError,
    ReportAttachmentTooLargeError,
    ReportAttachmentTooManyError,
    ReportAttachmentTotalTooLargeError,
)
from iris_agent.reports.attachments import AttachmentRepository
from iris_agent.reports.extraction import LocalAttachmentExtractor
from iris_agent.reports.errors import ReportValidationError


def test_package_exports_attachment_errors() -> None:
    from iris_agent.reports import ReportAttachmentExtractError, ReportAttachmentOcrUnavailableError

    assert issubclass(ReportAttachmentInvalidTypeError, ReportAttachmentError)
    assert issubclass(ReportAttachmentTooLargeError, ReportAttachmentError)
    assert issubclass(ReportAttachmentTooManyError, ReportAttachmentError)
    assert issubclass(ReportAttachmentTotalTooLargeError, ReportAttachmentError)
    assert issubclass(ReportAttachmentNotFoundError, ReportAttachmentError)
    assert issubclass(ReportAttachmentStorageError, ReportAttachmentError)
    assert issubclass(ReportAttachmentExtractError, ReportAttachmentError)
    assert issubclass(ReportAttachmentOcrUnavailableError, ReportAttachmentError)


def repository(tmp_path: Path, **limits: int) -> AttachmentRepository:
    return AttachmentRepository(
        tmp_path,
        max_file_bytes=limits.get("max_file_bytes", 20),
        max_total_bytes=limits.get("max_total_bytes", 40),
        max_count=limits.get("max_count", 2),
    )


class ExtractionFile:
    def __init__(self, name: str, content: bytes):
        self.name = name
        self.suffix = Path(name).suffix
        self._content = content

    def read_bytes(self) -> bytes:
        return self._content


def extraction_file(_tmp_path: Path, name: str, content: bytes) -> ExtractionFile:
    return ExtractionFile(name, content)


def test_extracts_utf8_text_and_reports_source(tmp_path: Path) -> None:
    result = LocalAttachmentExtractor(max_chars=100).extract(extraction_file(tmp_path, "notes.txt", "完成日报".encode()))
    assert result.text == "完成日报"
    assert result.sources[0].file_name == "notes.txt"


def test_extracts_markdown_as_utf8_text(tmp_path: Path) -> None:
    result = LocalAttachmentExtractor(max_chars=100).extract(extraction_file(tmp_path, "notes.md", b"# Done\n- report"))
    assert "# Done" in result.text


def test_extracts_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    from docx import Document
    document = Document()
    document.add_paragraph("段落内容")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "任务", "完成"
    path = tmp_path / "source.docx"
    document.save(path)
    result = LocalAttachmentExtractor(max_chars=100).extract(extraction_file(tmp_path, "source.docx", path.read_bytes()))
    assert "段落内容" in result.text
    assert "任务\t完成" in result.text


def test_extracts_pdf_text(tmp_path: Path) -> None:
    pdf = (b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
           b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n"
           b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n5 0 obj<</Length 37>>stream\nBT /F1 12 Tf 10 100 Td (PDF text) Tj ET\nendstream\nendobj\n"
           b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000241 00000 n \n0000000311 00000 n \ntrailer<</Size 6/Root 1 0 R>>\nstartxref\n397\n%%EOF")
    result = LocalAttachmentExtractor(max_chars=100).extract(extraction_file(tmp_path, "source.pdf", pdf))
    assert "PDF text" in result.text


def test_extracts_excel_sheet_names_and_limited_cells(tmp_path: Path) -> None:
    from openpyxl import Workbook
    workbook = Workbook()
    workbook.active.title = "进度"
    workbook.active.append(["任务", "状态"])
    workbook.active.append(["日报", "完成"])
    path = tmp_path / "source.xlsx"
    workbook.save(path)
    result = LocalAttachmentExtractor(max_chars=100).extract(extraction_file(tmp_path, "source.xlsx", path.read_bytes()))
    assert "工作表：进度" in result.text
    assert "日报\t完成" in result.text


def test_extraction_truncates_and_does_not_claim_empty_success(tmp_path: Path) -> None:
    extractor = LocalAttachmentExtractor(max_chars=4)
    assert extractor.extract(extraction_file(tmp_path, "notes.txt", b"abcdef")).text == "abcd"
    with pytest.raises(ReportAttachmentError) as error:
        extractor.extract(extraction_file(tmp_path, "empty.txt", b"   "))
    assert error.value.code == "report_attachment_extract_failed"


def test_extraction_wraps_parse_failure_and_reports_unavailable_ocr(tmp_path: Path) -> None:
    extractor = LocalAttachmentExtractor(max_chars=100)
    with pytest.raises(ReportAttachmentError) as corrupt:
        extractor.extract(extraction_file(tmp_path, "broken.pdf", b"not a PDF"))
    assert corrupt.value.code == "report_attachment_extract_failed"
    with pytest.raises(ReportAttachmentError) as ocr:
        extractor.extract(extraction_file(tmp_path, "picture.png", b"not-an-image"))
    assert ocr.value.code == "report_attachment_ocr_unavailable"


def test_save_normalizes_client_path_and_uses_server_generated_filename(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    saved = attachments.save(
        "2026-08-05", "../../secrets/notes.txt", b"safe", "text/plain", preserve=True
    )

    stored_path = attachments.path_for(saved.id)

    assert saved.original_name == "notes.txt"
    assert stored_path.parent == tmp_path / "2026-08-05" / "attachments"
    assert stored_path.name != "notes.txt"
    assert stored_path.suffix == ".txt"
    assert stored_path.read_bytes() == b"safe"


@pytest.mark.parametrize("original_name", ["/private/notes.txt", r"C:\\private\\notes.txt"])
def test_save_normalizes_posix_and_windows_absolute_client_paths(tmp_path: Path, original_name: str) -> None:
    attachments = repository(tmp_path)

    saved = attachments.save("2026-08-05", original_name, b"safe", "text/plain", preserve=True)

    assert saved.original_name == "notes.txt"
    assert attachments.path_for(saved.id).parent == tmp_path / "2026-08-05" / "attachments"


def test_save_allows_txt_and_lists_attachment_for_date(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    saved = attachments.save("2026-08-05", "daily.txt", b"hello", "text/plain", preserve=True)

    assert attachments.list_for_date("2026-08-05") == [saved]


def test_extraction_result_is_persisted_and_survives_restart(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    saved = attachments.save("2026-08-05", "daily.txt", b"hello", "text/plain", preserve=True)

    updated = attachments.set_extraction_result(
        saved.id,
        extraction_status="failed",
        extraction_message="无法提取日报附件文本",
    )

    assert updated.extraction_status == "failed"
    assert updated.extraction_message == "无法提取日报附件文本"
    assert repository(tmp_path).list_for_date("2026-08-05") == [updated]


def test_legacy_attachment_index_without_extraction_fields_derives_text_readiness(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    saved = attachments.save("2026-08-05", "daily.txt", b"hello", "text/plain", preserve=True)
    attachments.set_extraction_result(saved.id, extraction_status="ready", extracted_text="hello")
    index_path = tmp_path / "2026-08-05" / "attachments" / "index.json"
    payload = json.loads(index_path.read_text(encoding="utf-8"))
    del payload["attachments"][0]["extraction_status"]
    del payload["attachments"][0]["extraction_message"]
    index_path.write_text(json.dumps(payload), encoding="utf-8")

    restored = repository(tmp_path).list_for_date("2026-08-05")

    assert restored[0].extraction_status == "ready"
    assert restored[0].extraction_message is None


def test_save_rejects_unknown_extension(tmp_path: Path) -> None:
    with pytest.raises(ReportAttachmentError) as error:
        repository(tmp_path).save("2026-08-05", "malware.exe", b"no", "application/octet-stream", True)

    assert error.value.code == "report_attachment_invalid_type"


@pytest.mark.parametrize("media_type", ["image/png", None])
def test_save_rejects_mismatched_or_missing_media_type(tmp_path: Path, media_type: object) -> None:
    with pytest.raises(ReportAttachmentError) as error:
        repository(tmp_path).save("2026-08-05", "notes.txt", b"safe", media_type, True)  # type: ignore[arg-type]

    assert error.value.code == "report_attachment_invalid_type"


def test_save_rejects_empty_file(tmp_path: Path) -> None:
    with pytest.raises(ReportAttachmentError) as error:
        repository(tmp_path).save("2026-08-05", "empty.txt", b"", "text/plain", True)

    assert error.value.code == "report_attachment_invalid_type"


def test_save_rejects_file_larger_than_limit(tmp_path: Path) -> None:
    with pytest.raises(ReportAttachmentError) as error:
        repository(tmp_path, max_file_bytes=3).save("2026-08-05", "large.txt", b"four", "text/plain", True)

    assert error.value.code == "report_attachment_too_large"


def test_save_enforces_count_and_total_size_limits(tmp_path: Path) -> None:
    attachments = repository(tmp_path, max_total_bytes=5, max_count=1)
    attachments.save("2026-08-05", "first.txt", b"123", "text/plain", True)

    with pytest.raises(ReportAttachmentError) as count_error:
        attachments.save("2026-08-05", "second.txt", b"1", "text/plain", True)
    assert count_error.value.code == "report_attachment_too_many"

    by_size = repository(tmp_path / "by-size", max_total_bytes=5, max_count=2)
    by_size.save("2026-08-05", "first.txt", b"123", "text/plain", True)
    with pytest.raises(ReportAttachmentError) as total_error:
        by_size.save("2026-08-05", "second.txt", b"123", "text/plain", True)
    assert total_error.value.code == "report_attachment_total_too_large"


@pytest.mark.parametrize(
    ("max_count", "max_total_bytes"),
    [(1, 20), (4, 3)],
)
def test_concurrent_saves_do_not_exceed_count_or_total_limit(
    tmp_path: Path, max_count: int, max_total_bytes: int
) -> None:
    attachments = repository(tmp_path, max_count=max_count, max_total_bytes=max_total_bytes)
    start = threading.Barrier(8)
    outcomes: list[str] = []

    def save() -> None:
        start.wait()
        try:
            attachments.save("2026-08-05", "notes.txt", b"abc", "text/plain", True)
            outcomes.append("saved")
        except ReportAttachmentError:
            outcomes.append("rejected")

    workers = [threading.Thread(target=save) for _ in range(8)]
    for worker in workers:
        worker.start()
    for worker in workers:
        worker.join()

    saved = attachments.list_for_date("2026-08-05")
    assert len(saved) <= max_count
    assert sum(item.size_bytes for item in saved) <= max_total_bytes
    assert outcomes.count("saved") == len(saved)


def test_metadata_construction_failure_removes_written_file(tmp_path: Path, monkeypatch) -> None:
    attachments = repository(tmp_path)

    def fail_metadata(**_kwargs):
        raise ValueError("broken metadata")

    monkeypatch.setattr("iris_agent.reports.attachments.ReportAttachment", fail_metadata)

    with pytest.raises(ReportAttachmentStorageError):
        attachments.save("2026-08-05", "notes.txt", b"safe", "text/plain", True)

    assert not list((tmp_path / "2026-08-05" / "attachments").glob("*.txt"))
    assert attachments.list_for_date("2026-08-05") == []


def test_index_registration_failure_removes_written_file_and_quota_state(tmp_path: Path, monkeypatch) -> None:
    attachments = repository(tmp_path)

    def fail_index(*_args, **_kwargs):
        raise ReportAttachmentStorageError("index unavailable")

    monkeypatch.setattr(attachments, "_write_index", fail_index)

    with pytest.raises(ReportAttachmentStorageError):
        attachments.save("2026-08-05", "notes.txt", b"safe", "text/plain", True)

    assert not list((tmp_path / "2026-08-05" / "attachments").glob("*.txt"))
    assert attachments.list_for_date("2026-08-05") == []


def test_preserved_metadata_survives_restart_and_enforces_existing_quota(tmp_path: Path) -> None:
    attachments = repository(tmp_path, max_count=2, max_total_bytes=6)
    first = attachments.save("2026-08-05", "first.txt", b"abc", "text/plain", True)

    restarted = repository(tmp_path, max_count=2, max_total_bytes=6)

    assert restarted.list_for_date("2026-08-05") == [first]
    assert restarted.path_for(first.id).read_bytes() == b"abc"
    restarted.save("2026-08-05", "second.txt", b"abc", "text/plain", True)
    with pytest.raises(ReportAttachmentError) as error:
        restarted.save("2026-08-05", "third.txt", b"a", "text/plain", True)
    assert error.value.code == "report_attachment_too_many"


def test_restarted_repository_deletes_preserved_attachment_and_updates_index(tmp_path: Path) -> None:
    first = repository(tmp_path).save("2026-08-05", "first.txt", b"abc", "text/plain", True)
    restarted = repository(tmp_path)

    restarted.delete(first.id)

    assert restarted.list_for_date("2026-08-05") == []
    assert not (tmp_path / "2026-08-05" / "attachments" / f"{first.id}.txt").exists()
    assert repository(tmp_path).list_for_date("2026-08-05") == []


def test_separate_repositories_keep_concurrent_persisted_saves_and_quota(tmp_path: Path) -> None:
    first_repository = repository(tmp_path, max_count=2, max_total_bytes=6)
    second_repository = repository(tmp_path, max_count=2, max_total_bytes=6)
    start = threading.Barrier(2)
    failures: list[Exception] = []

    def save(target: AttachmentRepository, name: str) -> None:
        start.wait()
        try:
            target.save("2026-08-05", name, b"abc", "text/plain", True)
        except Exception as exc:  # pragma: no cover - asserted below
            failures.append(exc)

    first = threading.Thread(target=save, args=(first_repository, "first.txt"))
    second = threading.Thread(target=save, args=(second_repository, "second.txt"))
    first.start()
    second.start()
    first.join()
    second.join()

    assert failures == []
    restarted = repository(tmp_path, max_count=2, max_total_bytes=6)
    assert {item.original_name for item in restarted.list_for_date("2026-08-05")} == {"first.txt", "second.txt"}
    with pytest.raises(ReportAttachmentError) as error:
        restarted.save("2026-08-05", "third.txt", b"a", "text/plain", True)
    assert error.value.code == "report_attachment_too_many"


def _make_directory_symlink(link: Path, target: Path) -> None:
    try:
        os.symlink(target, link, target_is_directory=True)
    except OSError as exc:
        pytest.skip(f"symlink creation is unavailable: {exc}")


def test_rejects_external_attachments_directory_symlink_during_load(tmp_path: Path) -> None:
    date_directory = tmp_path / "2026-08-05"
    date_directory.mkdir()
    external = tmp_path / "external"
    external.mkdir()
    _make_directory_symlink(date_directory / "attachments", external)

    with pytest.raises(ReportAttachmentStorageError):
        repository(tmp_path)


def test_delete_refuses_external_attachments_directory_symlink(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    saved = attachments.save("2026-08-05", "safe.txt", b"safe", "text/plain", True)
    directory = tmp_path / "2026-08-05" / "attachments"
    external = tmp_path / "external"
    external.mkdir()
    external_file = external / f"{saved.id}.txt"
    external_file.write_bytes(b"outside")
    shutil.rmtree(directory)
    _make_directory_symlink(directory, external)

    with pytest.raises(ReportAttachmentStorageError):
        attachments.delete(saved.id)

    assert external_file.read_bytes() == b"outside"


@pytest.mark.parametrize(
    "bad_field",
    [
        {"preserve": "true"},
        {"size_bytes": True},
        {"created_at": "now"},
        {"extracted_text": 3},
    ],
)
def test_rejects_index_records_with_wrong_json_types(tmp_path: Path, bad_field: dict[str, object]) -> None:
    directory = tmp_path / "2026-08-05" / "attachments"
    directory.mkdir(parents=True)
    record = {
        "id": "0f0f0f0f-0f0f-4f0f-8f0f-0f0f0f0f0f0f",
        "original_name": "safe.txt",
        "media_type": "text/plain",
        "size_bytes": 4,
        "preserve": True,
        "status": "preserved",
        "extracted_text": None,
        "created_at": 1.0,
        "file_name": "0f0f0f0f-0f0f-4f0f-8f0f-0f0f0f0f0f0f.txt",
    }
    record.update(bad_field)
    size = 1 if bad_field == {"size_bytes": True} else 4
    (directory / record["file_name"]).write_bytes(b"x" * size)
    (directory / "index.json").write_text(json.dumps({"attachments": [record]}), encoding="utf-8")

    with pytest.raises(ReportAttachmentStorageError):
        repository(tmp_path)


def test_rejects_attachment_file_symlink_that_escapes_storage(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    saved = attachments.save("2026-08-05", "safe.txt", b"safe", "text/plain", True)
    external = tmp_path / "outside.txt"
    external.write_bytes(b"outside")
    stored = tmp_path / "2026-08-05" / "attachments" / f"{saved.id}.txt"
    stored.unlink()
    try:
        os.symlink(external, stored)
    except OSError as exc:
        pytest.skip(f"symlink creation is unavailable: {exc}")

    with pytest.raises(ReportAttachmentStorageError):
        attachments.path_for(saved.id)
    with pytest.raises(ReportAttachmentStorageError):
        attachments.delete(saved.id)
    assert external.read_bytes() == b"outside"


def test_separate_repositories_enforce_temporary_attachment_quota(tmp_path: Path) -> None:
    first_repository = repository(tmp_path, max_count=1)
    second_repository = repository(tmp_path, max_count=1)
    start = threading.Barrier(2)
    outcomes: list[str] = []

    def save(target: AttachmentRepository) -> None:
        start.wait()
        try:
            target.save("2026-08-05", "temp.txt", b"abc", "text/plain", False)
            outcomes.append("saved")
        except ReportAttachmentError:
            outcomes.append("rejected")

    first = threading.Thread(target=save, args=(first_repository,))
    second = threading.Thread(target=save, args=(second_repository,))
    first.start()
    second.start()
    first.join()
    second.join()

    restarted = repository(tmp_path, max_count=1)
    assert outcomes.count("saved") == 1
    assert len(restarted.list_for_date("2026-08-05")) == 1


def test_expired_dead_process_lock_is_reclaimed(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    lock_path = tmp_path / "2026-08-05" / ".attachments.lock"
    lock_path.parent.mkdir()
    lock_path.write_text(json.dumps({"pid": 999_999, "created_at": 0.0, "token": "stale"}), encoding="utf-8")

    attachments.save("2026-08-05", "safe.txt", b"safe", "text/plain", True)

    assert not lock_path.exists()


def test_active_process_lock_is_not_reclaimed(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    lock_path = tmp_path / "2026-08-05" / ".attachments.lock"
    lock_path.parent.mkdir()
    lock_path.write_text(
        json.dumps({"pid": os.getpid(), "created_at": 0.0, "token": "active"}), encoding="utf-8"
    )

    assert attachments._reclaim_stale_lock(lock_path) is False
    assert lock_path.exists()


def test_expired_empty_lock_is_reclaimed_by_mtime(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    lock_path = tmp_path / "2026-08-05" / ".attachments.lock"
    lock_path.parent.mkdir()
    lock_path.write_bytes(b"")
    os.utime(lock_path, (0, 0))

    attachments.save("2026-08-05", "safe.txt", b"safe", "text/plain", True)

    assert not lock_path.exists()


def test_save_rejects_directory_that_changes_after_initial_validation(tmp_path: Path, monkeypatch) -> None:
    attachments = repository(tmp_path)
    calls = 0
    real_directory_for = attachments._directory_for

    def change_after_first_check(*args, **kwargs):
        nonlocal calls
        calls += 1
        if calls == 2:
            raise ReportAttachmentStorageError("directory changed")
        return real_directory_for(*args, **kwargs)

    monkeypatch.setattr(attachments, "_directory_for", change_after_first_check)

    with pytest.raises(ReportAttachmentStorageError):
        attachments.save("2026-08-05", "safe.txt", b"safe", "text/plain", True)


def test_report_attachment_rejects_non_finite_created_at() -> None:
    from iris_agent.reports.attachments import ReportAttachment

    with pytest.raises(ValueError):
        ReportAttachment(
            id="0f0f0f0f-0f0f-4f0f-8f0f-0f0f0f0f0f0f",
            original_name="safe.txt",
            media_type="text/plain",
            size_bytes=1,
            preserve=True,
            status="preserved",
            created_at=math.nan,
        )


def test_cleanup_removes_only_temporary_attachments(tmp_path: Path) -> None:
    attachments = repository(tmp_path)
    temporary = attachments.save("2026-08-05", "temp.txt", b"temp", "text/plain", preserve=False)
    preserved = attachments.save("2026-08-05", "keep.txt", b"keep", "text/plain", preserve=True)

    attachments.cleanup([temporary.id, preserved.id])

    with pytest.raises(ReportAttachmentError) as missing:
        attachments.path_for(temporary.id)
    assert missing.value.code == "report_attachment_not_found"
    assert attachments.path_for(preserved.id).exists()


def test_repository_rejects_invalid_date(tmp_path: Path) -> None:
    with pytest.raises(ReportValidationError) as error:
        repository(tmp_path).save("../../secret", "daily.txt", b"safe", "text/plain", True)

    assert error.value.code == "report_invalid_date"

