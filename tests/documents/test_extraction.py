from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import sys
from types import SimpleNamespace

import pytest

from iris_agent.documents import DocumentExtractFailedError, LocalDocumentExtractor


@dataclass
class SourceFile:
    name: str
    content: bytes

    @property
    def suffix(self) -> str:
        return Path(self.name).suffix

    def read_bytes(self) -> bytes:
        return self.content


def test_extracts_txt_and_markdown_with_source_and_global_truncation() -> None:
    extractor = LocalDocumentExtractor(max_text_chars=4)

    text = extractor.extract(SourceFile("notes.txt", "abcdef".encode()))
    markdown = LocalDocumentExtractor(max_text_chars=100).extract(SourceFile("notes.md", b"# Title\n- item"))

    assert text.text == "abcd"
    assert text.truncated is True
    assert text.sources[0].file_name == "notes.txt"
    assert "# Title" in markdown.text


def test_extracts_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("段落内容")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "任务"
    table.cell(0, 1).text = "完成"
    path = tmp_path / "source.docx"
    document.save(path)

    result = LocalDocumentExtractor(max_text_chars=100).extract(SourceFile("source.docx", path.read_bytes()))

    assert "段落内容" in result.text
    assert "任务\t完成" in result.text
    assert {source.location for source in result.sources} >= {"段落", "表格 1"}


def test_extracts_pdf_pages_with_page_sources() -> None:
    pdf = (
        b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n"
        b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n5 0 obj<</Length 37>>stream\nBT /F1 12 Tf 10 100 Td (PDF text) Tj ET\nendstream\nendobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000241 00000 n \n0000000311 00000 n \ntrailer<</Size 6/Root 1 0 R>>\nstartxref\n397\n%%EOF"
    )

    result = LocalDocumentExtractor(max_text_chars=100).extract(SourceFile("source.pdf", pdf))

    assert "PDF text" in result.text
    assert result.sources[0].location == "第 1 页"


def test_extracts_xlsx_rows_and_limits_rows_and_columns(tmp_path: Path) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "进度"
    sheet.append(["任务", "状态", "忽略列"] + ["x"] * 18)
    sheet.append(["日报", "完成"])
    path = tmp_path / "source.xlsx"
    workbook.save(path)

    result = LocalDocumentExtractor(max_text_chars=10_000).extract(SourceFile("source.xlsx", path.read_bytes()))

    assert "工作表：进度" in result.text
    assert "日报\t完成" in result.text
    assert result.sources[0].location == "工作表：进度"


def test_extracts_xls_rows_and_limits_rows_and_columns(monkeypatch: pytest.MonkeyPatch) -> None:
    class Sheet:
        name = "历史"
        nrows = 101
        ncols = 21

        @staticmethod
        def cell_value(row: int, column: int) -> str:
            return f"{row}:{column}"

    workbook = SimpleNamespace(sheets=lambda: [Sheet()], release_resources=lambda: None)
    monkeypatch.setitem(sys.modules, "xlrd", SimpleNamespace(open_workbook=lambda **_kwargs: workbook))

    result = LocalDocumentExtractor(max_text_chars=20_000).extract(SourceFile("source.xls", b"legacy-xls"))

    assert "工作表：历史" in result.text
    assert "0:0" in result.text
    assert "99:19" in result.text
    assert "100:0" not in result.text
    assert "0:20" not in result.text


@pytest.mark.parametrize("name, content", [("broken.pdf", b"not a pdf"), ("empty.txt", b"  "), ("image.png", b"png")])
def test_rejects_unextractable_or_unsupported_content(name: str, content: bytes) -> None:
    with pytest.raises(DocumentExtractFailedError) as error:
        LocalDocumentExtractor(max_text_chars=100).extract(SourceFile(name, content))

    assert error.value.code == "document_extract_failed"
